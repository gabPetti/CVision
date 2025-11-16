"""
CVision - Shared Library
Componentes reutilizáveis: ResponseBuilder, LLMProvider, CVTextExtractor
"""

import os
import io
import logging
from typing import Optional, Dict, Any
from pathlib import Path

import PyPDF2
from docx import Document

logger = logging.getLogger(__name__)


# ============================================================================
# CV TEXT EXTRACTOR - Parser para PDF e DOCX
# ============================================================================

class CVTextExtractor:
    """Extrai texto de arquivos CV (PDF, DOCX, TXT)"""

    SUPPORTED_FORMATS = {'.pdf', '.docx', '.doc', '.txt'}
    MAX_FILE_SIZE = 10 * 1024 * 1024  # 10MB

    @staticmethod
    def is_valid_file(file_path: str | Path) -> bool:
        """Valida se arquivo é suportado"""
        path = Path(file_path)
        return path.suffix.lower() in CVTextExtractor.SUPPORTED_FORMATS

    @staticmethod
    def extract_text(file_path: str | Path, file_content: Optional[bytes] = None) -> str:
        """
        Extrai texto do arquivo CV

        Args:
            file_path: Caminho do arquivo ou nome
            file_content: Conteúdo do arquivo em bytes (opcional)

        Returns:
            Texto extraído do CV
        """
        path = Path(file_path)
        suffix = path.suffix.lower()

        if not CVTextExtractor.is_valid_file(path):
            raise ValueError(f"Formato não suportado: {suffix}")

        try:
            if suffix == '.pdf':
                return CVTextExtractor._extract_pdf(file_path, file_content)
            elif suffix in {'.docx', '.doc'}:
                return CVTextExtractor._extract_docx(file_path, file_content)
            elif suffix == '.txt':
                return CVTextExtractor._extract_txt(file_path, file_content)
        except Exception as e:
            logger.error(f"Erro ao extrair texto: {e}")
            raise

    @staticmethod
    def _extract_pdf(file_path: str | Path, file_content: Optional[bytes] = None) -> str:
        """Extrai texto de PDF"""
        try:
            if file_content:
                pdf_reader = PyPDF2.PdfReader(io.BytesIO(file_content))
            else:
                with open(file_path, 'rb') as f:
                    pdf_reader = PyPDF2.PdfReader(f)

            text = ""
            for page in pdf_reader.pages:
                text += page.extract_text() + "\n"

            return text.strip()
        except Exception as e:
            logger.error(f"Erro ao extrair PDF: {e}")
            raise

    @staticmethod
    def _extract_docx(file_path: str | Path, file_content: Optional[bytes] = None) -> str:
        """Extrai texto de DOCX"""
        try:
            if file_content:
                doc = Document(io.BytesIO(file_content))
            else:
                doc = Document(file_path)

            text = ""
            for paragraph in doc.paragraphs:
                text += paragraph.text + "\n"

            return text.strip()
        except Exception as e:
            logger.error(f"Erro ao extrair DOCX: {e}")
            raise

    @staticmethod
    def _extract_txt(file_path: str | Path, file_content: Optional[bytes] = None) -> str:
        """Extrai texto de TXT"""
        try:
            if file_content:
                return file_content.decode('utf-8')
            else:
                with open(file_path, 'r', encoding='utf-8') as f:
                    return f.read()
        except Exception as e:
            logger.error(f"Erro ao extrair TXT: {e}")
            raise


# ============================================================================
# RESPONSE BUILDER - Standardiza respostas HTTP
# ============================================================================

class ResponseBuilder:
    """Constrói respostas padronizadas para a API"""

    @staticmethod
    def success(data: Any, status_code: int = 200) -> tuple[dict, int]:
        """Resposta de sucesso"""
        return data, status_code

    @staticmethod
    def error(message: str, status_code: int = 400, details: Optional[dict] = None) -> tuple[dict, int]:
        """Resposta de erro"""
        return {
            'message': message,
            'status': status_code
        }, status_code

    @staticmethod
    def validation_error(errors: Dict[str, str], status_code: int = 422) -> tuple[dict, int]:
        """Resposta de erro de validação"""
        return {
            'success': False,
            'message': 'Erro de validação',
            'errors': errors,
            'status': status_code
        }, status_code


# ============================================================================
# LLM PROVIDER - Singleton Pattern
# ============================================================================

class LLMProvider:
    """Singleton para gerenciar instância do LLM (Gemini 2.5 Flash)"""

    _instance: Optional['LLMProvider'] = None
    _llm: Optional[Any] = None

    def __new__(cls):
        if cls._instance is None:
            cls._instance = super().__new__(cls)
        return cls._instance

    def get_llm(self):
        """Obtém instância do LLM com lazy loading"""
        if self._llm is None:
            from langchain_google_genai import ChatGoogleGenerativeAI

            api_key = os.getenv('GEMINI_API_KEY')
            if not api_key:
                raise ValueError(
                    "GEMINI_API_KEY não configurada. "
                    "Configure em .env ou variável de ambiente."
                )

            logger.info("Inicializando ChatGoogleGenerativeAI...")
            self._llm = ChatGoogleGenerativeAI(
                model="gemini-2.5-flash",
                api_key=api_key,
                temperature=0
            )

        return self._llm
